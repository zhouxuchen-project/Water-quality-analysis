from __future__ import annotations

from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from scipy.stats import spearmanr


ROOT = Path(__file__).resolve().parents[1]
PROCESSED = ROOT / "data" / "processed"
SENSITIVITY = ROOT / "results" / "7.5_sensitivity"
THESIS_RESULTS = ROOT / "results" / "thesis"
RESULTS = ROOT / "results" / "7.6_site_profiles"

UNCENSORED_SAMPLES = SENSITIVITY / "7.5_Uncensored_PbZn_Flow_Samples.csv"
MANUAL_SAMPLES = PROCESSED / "pbzn_manual_flow_validation_samples.csv"
ROBUST_STATIONS = THESIS_RESULTS / "Table_4_4_Robust_Candidate_Stations.csv"
THRESHOLDS = SENSITIVITY / "7.5_Censoring_Thresholds.csv"

LOW_FLOW_PERCENTILE = 0.25
VERY_LOW_FLOW_PERCENTILE = 0.10

SYSTEMS = {
    "S83019": "Esgair Mwyn / Nant y Garw",
    "S83018": "Esgair Mwyn / Nant y Garw",
    "S83017": "Esgair Mwyn / Nant y Garw",
    "S83020": "Esgair Mwyn / Nant y Garw",
    "S83021": "Esgair Mwyn / Nant y Garw",
    "S35767": "Nant y Watcyn",
    "S35279": "Esgair Hir / Eastern Ffraith",
    "S6320066": "Esgair Hir / Eastern Ffraith",
    "S35582": "Wemyss",
}

SYSTEM_ORDER = [
    "Esgair Mwyn / Nant y Garw",
    "Nant y Watcyn",
    "Esgair Hir / Eastern Ffraith",
    "Wemyss",
]

SYSTEM_SHORT = {
    "Esgair Mwyn / Nant y Garw": "Esgair Mwyn",
    "Nant y Watcyn": "Nant y Watcyn",
    "Esgair Hir / Eastern Ffraith": "Esgair Hir / E. Ffraith",
    "Wemyss": "Wemyss",
}

PALETTE = {
    "Esgair Mwyn / Nant y Garw": "#356B8C",
    "Nant y Watcyn": "#2F7D61",
    "Esgair Hir / Eastern Ffraith": "#B56B38",
    "Wemyss": "#8A4F73",
}

DRIVER_ORDER = [
    "High Pb and low Zn",
    "High Pb only",
    "Low Zn only",
    "Neither marginal threshold",
]

DRIVER_COLOURS = {
    "High Pb and low Zn": "#A93C3C",
    "High Pb only": "#D77A36",
    "Low Zn only": "#3E7DA6",
    "Neither marginal threshold": "#9A9A9A",
}


def truthy(series: pd.Series) -> pd.Series:
    return series.astype(str).str.lower().isin({"true", "1", "yes"})


def safe_ratio(numerator: float, denominator: float) -> float:
    if pd.isna(numerator) or pd.isna(denominator) or denominator == 0:
        return np.nan
    return float(numerator / denominator)


def safe_spearman(x: pd.Series, y: pd.Series, minimum: int = 5) -> tuple[int, float, float]:
    frame = pd.DataFrame({"x": x, "y": y}).replace([np.inf, -np.inf], np.nan).dropna()
    if len(frame) < minimum or frame["x"].nunique() < 2 or frame["y"].nunique() < 2:
        return len(frame), np.nan, np.nan
    result = spearmanr(frame["x"], frame["y"])
    return len(frame), float(result.statistic), float(result.pvalue)


def benjamini_hochberg(p_values: pd.Series) -> pd.Series:
    adjusted = pd.Series(np.nan, index=p_values.index, dtype=float)
    valid = p_values.dropna().sort_values()
    if valid.empty:
        return adjusted
    count = len(valid)
    raw = valid.to_numpy(float) * count / np.arange(1, count + 1)
    corrected = np.minimum.accumulate(raw[::-1])[::-1]
    adjusted.loc[valid.index] = np.minimum(corrected, 1.0)
    return adjusted


def season_for_month(month: int) -> str:
    if month in (12, 1, 2):
        return "Winter"
    if month in (3, 4, 5):
        return "Spring"
    if month in (6, 7, 8):
        return "Summer"
    return "Autumn"


def driver_category(row: pd.Series) -> str:
    high_pb = bool(row["high_lead"])
    low_zn = bool(row["low_zinc"])
    if high_pb and low_zn:
        return "High Pb and low Zn"
    if high_pb:
        return "High Pb only"
    if low_zn:
        return "Low Zn only"
    return "Neither marginal threshold"


def setup_plot_style() -> None:
    sns.set_theme(style="whitegrid", context="paper")
    plt.rcParams.update(
        {
            "font.family": "Arial",
            "font.size": 9.5,
            "axes.titlesize": 11,
            "axes.labelsize": 10,
            "legend.fontsize": 8,
            "figure.dpi": 130,
            "savefig.dpi": 300,
            "savefig.bbox": "tight",
            "axes.spines.top": False,
            "axes.spines.right": False,
        }
    )


def load_analysis_samples() -> tuple[pd.DataFrame, pd.DataFrame, dict[str, float]]:
    robust = pd.read_csv(ROBUST_STATIONS)
    robust["station_id"] = robust["station_id"].astype(str)
    robust = robust[robust["station_id"].isin(SYSTEMS)].copy()
    if set(robust["station_id"]) != set(SYSTEMS):
        missing = sorted(set(SYSTEMS) - set(robust["station_id"]))
        raise ValueError(f"Robust station table is missing: {missing}")

    samples = pd.read_csv(UNCENSORED_SAMPLES)
    samples["station_id"] = samples["station_id"].astype(str)
    samples = samples[samples["station_id"].isin(SYSTEMS)].copy()
    samples["sample_date"] = pd.to_datetime(samples["sample_date"])

    manual = pd.read_csv(MANUAL_SAMPLES)
    manual["station_id"] = manual["station_id"].astype(str)
    manual = manual[manual["station_id"].isin(SYSTEMS)].copy()
    manual["sample_date"] = pd.to_datetime(manual["sample_date"])
    manual_columns = [
        "station_id",
        "sample_date",
        "calcium",
        "copper",
        "hardness",
        "ph",
        "manual_nrfa_id",
        "manual_nrfa_name",
        "manual_nrfa_river",
        "manual_match_distance_km",
        "manual_decision",
        "manual_confidence",
        "manual_flow_date",
        "manual_flow_m3s",
        "manual_flow_flag",
        "manual_flow_percentile",
        "manual_flow_p10_m3s",
        "manual_flow_p25_m3s",
        "manual_flow_p50_m3s",
        "manual_has_flow_match",
    ]
    manual = manual[manual_columns].drop_duplicates(["station_id", "sample_date"])
    samples = samples.merge(
        manual,
        on=["station_id", "sample_date"],
        how="left",
        validate="one_to_one",
    )

    for column in [
        "lead_value",
        "zinc_value",
        "pb_zn_ratio",
        "calcium",
        "copper",
        "ph",
        "manual_flow_m3s",
        "manual_flow_percentile",
    ]:
        samples[column] = pd.to_numeric(samples[column], errors="coerce")
    for column in ["high_lead", "low_zinc", "high_pb_low_zinc", "high_pb_zn_ratio"]:
        samples[column] = truthy(samples[column])

    samples["system"] = samples["station_id"].map(SYSTEMS)
    samples["system_short"] = samples["system"].map(SYSTEM_SHORT)
    samples["sample_day"] = samples["sample_date"].dt.floor("D")
    samples["year"] = samples["sample_date"].dt.year
    samples["month"] = samples["sample_date"].dt.month
    samples["season"] = samples["month"].map(season_for_month)
    samples["manual_has_flow"] = samples["manual_flow_percentile"].notna()
    samples["manual_low_flow"] = (
        samples["manual_has_flow"]
        & samples["manual_flow_percentile"].le(LOW_FLOW_PERCENTILE)
    )
    samples["manual_very_low_flow"] = (
        samples["manual_has_flow"]
        & samples["manual_flow_percentile"].le(VERY_LOW_FLOW_PERCENTILE)
    )
    samples["manual_flow_class"] = np.select(
        [
            samples["manual_very_low_flow"],
            samples["manual_low_flow"],
            samples["manual_has_flow"],
        ],
        ["Very low flow", "Low flow", "Normal/high flow"],
        default="No matched flow",
    )
    samples["primary_event"] = samples["manual_low_flow"] & samples["high_pb_zn_ratio"]
    samples["strict_event"] = samples["manual_low_flow"] & samples["high_pb_low_zinc"]
    samples["event_driver"] = samples.apply(driver_category, axis=1)
    samples.loc[~samples["primary_event"], "event_driver"] = "Not a primary event"

    normal_flow = samples["manual_flow_percentile"].gt(LOW_FLOW_PERCENTILE)
    normal_medians = samples.loc[normal_flow].groupby("station_id").agg(
        station_normal_pb_ug_l=("lead_value", "median"),
        station_normal_zn_ug_l=("zinc_value", "median"),
        station_normal_pb_zn_ratio=("pb_zn_ratio", "median"),
    )
    samples = samples.merge(normal_medians, on="station_id", how="left", validate="many_to_one")
    samples["pb_relative_to_station_normal"] = (
        samples["lead_value"] / samples["station_normal_pb_ug_l"]
    )
    samples["zn_relative_to_station_normal"] = (
        samples["zinc_value"] / samples["station_normal_zn_ug_l"]
    )
    samples["ratio_relative_to_station_normal"] = (
        samples["pb_zn_ratio"] / samples["station_normal_pb_zn_ratio"]
    )

    thresholds_frame = pd.read_csv(THRESHOLDS)
    threshold_row = thresholds_frame.loc[
        thresholds_frame["scenario"].eq("uncensored_only")
    ].iloc[0]
    thresholds = {
        "lead_high_ug_l": float(threshold_row["lead_high_ug_l"]),
        "zinc_low_ug_l": float(threshold_row["zinc_low_ug_l"]),
        "high_pb_zn_ratio": float(threshold_row["high_pb_zn_ratio"]),
        "low_flow_percentile": LOW_FLOW_PERCENTILE,
    }

    ordered_columns = [
        "system",
        "station_id",
        "station_name",
        "sample_date",
        "lead_value",
        "zinc_value",
        "pb_zn_ratio",
        "high_lead",
        "low_zinc",
        "high_pb_low_zinc",
        "high_pb_zn_ratio",
        "calcium",
        "copper",
        "ph",
        "manual_nrfa_id",
        "manual_nrfa_name",
        "manual_nrfa_river",
        "manual_match_distance_km",
        "manual_flow_date",
        "manual_flow_m3s",
        "manual_flow_percentile",
        "manual_flow_class",
        "manual_low_flow",
        "manual_very_low_flow",
        "primary_event",
        "strict_event",
        "event_driver",
        "station_normal_pb_ug_l",
        "station_normal_zn_ug_l",
        "station_normal_pb_zn_ratio",
        "pb_relative_to_station_normal",
        "zn_relative_to_station_normal",
        "ratio_relative_to_station_normal",
        "season",
        "year",
        "station_type",
        "wfd_c2_mgt_catchment_name",
    ]
    samples = samples[ordered_columns].sort_values(["system", "station_id", "sample_date"])
    return samples, robust, thresholds


def summarize_group(group: pd.DataFrame) -> dict[str, float | int | str]:
    flow = group[group["manual_flow_percentile"].notna()]
    events = flow[flow["primary_event"]]
    normal = flow[flow["manual_flow_percentile"] > LOW_FLOW_PERCENTILE]
    low_flow = flow[flow["manual_low_flow"]]
    return {
        "n_uncensored_pbzn": len(group),
        "first_sample_date": group["sample_date"].min().date(),
        "last_sample_date": group["sample_date"].max().date(),
        "n_with_manual_flow": len(flow),
        "n_low_flow_samples": len(low_flow),
        "n_low_flow_high_ratio_events": int(events.shape[0]),
        "n_strict_high_pb_low_zn_events": int(events["strict_event"].sum()),
        "event_pct_of_flow_samples": safe_ratio(len(events) * 100, len(flow)),
        "event_pct_of_low_flow_samples": safe_ratio(len(events) * 100, len(low_flow)),
        "median_pb_ug_l": group["lead_value"].median(),
        "median_zn_ug_l": group["zinc_value"].median(),
        "median_pb_zn_ratio": group["pb_zn_ratio"].median(),
        "maximum_pb_zn_ratio": group["pb_zn_ratio"].max(),
        "event_median_pb_ug_l": events["lead_value"].median(),
        "normal_flow_median_pb_ug_l": normal["lead_value"].median(),
        "event_to_normal_pb_fold": safe_ratio(
            events["lead_value"].median(), normal["lead_value"].median()
        ),
        "event_median_zn_ug_l": events["zinc_value"].median(),
        "normal_flow_median_zn_ug_l": normal["zinc_value"].median(),
        "event_to_normal_zn_fold": safe_ratio(
            events["zinc_value"].median(), normal["zinc_value"].median()
        ),
        "event_median_pb_zn_ratio": events["pb_zn_ratio"].median(),
        "normal_flow_median_pb_zn_ratio": normal["pb_zn_ratio"].median(),
        "event_to_normal_ratio_fold": safe_ratio(
            events["pb_zn_ratio"].median(), normal["pb_zn_ratio"].median()
        ),
        "n_ph": int(group["ph"].notna().sum()),
        "median_ph": group["ph"].median(),
        "event_median_ph": events["ph"].median(),
        "normal_flow_median_ph": normal["ph"].median(),
        "n_calcium": int(group["calcium"].notna().sum()),
        "median_calcium_mg_l": group["calcium"].median(),
        "event_median_calcium_mg_l": events["calcium"].median(),
        "normal_flow_median_calcium_mg_l": normal["calcium"].median(),
    }


def make_site_summary(samples: pd.DataFrame, robust: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for station_id, group in samples.groupby("station_id", sort=False):
        row = {
            "system": SYSTEMS[station_id],
            "station_id": station_id,
            "station_name": group["station_name"].iloc[0],
            "manual_match_distance_km": group["manual_match_distance_km"].dropna().iloc[0]
            if group["manual_match_distance_km"].notna().any()
            else np.nan,
            **summarize_group(group),
        }
        n_flow, rho_flow, p_flow = safe_spearman(
            group["manual_flow_percentile"], group["pb_zn_ratio"]
        )
        row.update(
            {
                "n_ratio_flow_correlation": n_flow,
                "spearman_ratio_vs_flow_rho": rho_flow,
                "spearman_ratio_vs_flow_p": p_flow,
            }
        )
        rows.append(row)
    site = pd.DataFrame(rows)
    station_metadata = robust[
        [
            "station_id",
            "station_type",
            "easting",
            "northing",
            "manual_nrfa_id",
            "manual_nrfa_name",
            "manual_decision",
        ]
    ].copy()
    site = site.merge(station_metadata, on="station_id", how="left")
    site["spearman_ratio_vs_flow_fdr_q"] = benjamini_hochberg(
        site["spearman_ratio_vs_flow_p"]
    )
    site["flow_hypothesis_interpretation"] = "Co-occurrence without clear low-flow amplification"
    site.loc[
        site["event_to_normal_ratio_fold"].ge(1.25)
        & site["spearman_ratio_vs_flow_fdr_q"].ge(0.05),
        "flow_hypothesis_interpretation",
    ] = "Suggestive event enrichment; monotonic test not significant"
    site.loc[
        site["spearman_ratio_vs_flow_fdr_q"].lt(0.05)
        & site["spearman_ratio_vs_flow_rho"].lt(0)
        & site["event_to_normal_ratio_fold"].ge(1.25),
        "flow_hypothesis_interpretation",
    ] = "Low-flow amplification supported"
    site.loc[
        site["spearman_ratio_vs_flow_fdr_q"].lt(0.05)
        & site["spearman_ratio_vs_flow_rho"].gt(0),
        "flow_hypothesis_interpretation",
    ] = "Significant association opposite to low-flow hypothesis"
    site["system_order"] = site["system"].map({name: i for i, name in enumerate(SYSTEM_ORDER)})
    site = site.sort_values(
        ["system_order", "n_low_flow_high_ratio_events", "station_id"],
        ascending=[True, False, True],
    ).drop(columns="system_order")
    return site


def make_system_summary(samples: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for system in SYSTEM_ORDER:
        group = samples[samples["system"].eq(system)]
        row = {"system": system, "n_stations": group["station_id"].nunique(), **summarize_group(group)}
        events = group[group["primary_event"]]
        row.update(
            {
                "event_station_normalised_pb_median": events[
                    "pb_relative_to_station_normal"
                ].median(),
                "event_station_normalised_zn_median": events[
                    "zn_relative_to_station_normal"
                ].median(),
                "event_station_normalised_ratio_median": events[
                    "ratio_relative_to_station_normal"
                ].median(),
            }
        )
        n_flow, rho_flow, p_flow = safe_spearman(
            group["manual_flow_percentile"], group["pb_zn_ratio"]
        )
        row.update(
            {
                "n_ratio_flow_correlation": n_flow,
                "spearman_ratio_vs_flow_rho": rho_flow,
                "spearman_ratio_vs_flow_p": p_flow,
            }
        )
        rows.append(row)
    return pd.DataFrame(rows)


def make_driver_summary(samples: pd.DataFrame) -> pd.DataFrame:
    events = samples[samples["primary_event"]].copy()
    rows = []
    for system in [*SYSTEM_ORDER, "All four systems"]:
        group = events if system == "All four systems" else events[events["system"].eq(system)]
        for category in DRIVER_ORDER:
            count = int(group["event_driver"].eq(category).sum())
            rows.append(
                {
                    "system": system,
                    "driver_category": category,
                    "n_events": count,
                    "pct_of_system_events": safe_ratio(count * 100, len(group)),
                }
            )
    return pd.DataFrame(rows)


def make_chemistry_context(samples: pd.DataFrame) -> pd.DataFrame:
    rows = []
    groupings: Iterable[tuple[str, pd.DataFrame]] = [
        ("All four systems", samples),
        *[(system, samples[samples["system"].eq(system)]) for system in SYSTEM_ORDER],
        *[
            (f"Station {station_id}", group)
            for station_id, group in samples.groupby("station_id", sort=True)
        ],
    ]
    for label, group in groupings:
        events = group[group["primary_event"]]
        normal = group[
            group["manual_flow_percentile"].notna()
            & group["manual_flow_percentile"].gt(LOW_FLOW_PERCENTILE)
        ]
        for variable, unit in [
            ("manual_flow_percentile", "fraction"),
            ("ph", "pH units"),
            ("calcium", "mg/L"),
            ("lead_value", "ug/L"),
            ("zinc_value", "ug/L"),
        ]:
            n, rho, p_value = safe_spearman(group[variable], group["pb_zn_ratio"])
            rows.append(
                {
                    "analysis_group": label,
                    "variable": variable,
                    "unit": unit,
                    "n_complete_pairs": n,
                    "spearman_rho_with_pb_zn_ratio": rho,
                    "spearman_p_value": p_value,
                    "all_sample_median": group[variable].median(),
                    "event_median": events[variable].median(),
                    "normal_flow_median": normal[variable].median(),
                    "event_to_normal_fold": safe_ratio(
                        events[variable].median(), normal[variable].median()
                    )
                    if variable != "ph"
                    else np.nan,
                    "event_minus_normal": (
                        events[variable].median() - normal[variable].median()
                    ),
                }
            )
    return pd.DataFrame(rows)


def make_season_summary(samples: pd.DataFrame) -> pd.DataFrame:
    rows = []
    seasons = ["Winter", "Spring", "Summer", "Autumn"]
    for system in [*SYSTEM_ORDER, "All four systems"]:
        base = samples if system == "All four systems" else samples[samples["system"].eq(system)]
        for season in seasons:
            group = base[base["season"].eq(season)]
            flow = group[group["manual_flow_percentile"].notna()]
            low = flow[flow["manual_low_flow"]]
            events = flow[flow["primary_event"]]
            rows.append(
                {
                    "system": system,
                    "season": season,
                    "n_samples": len(group),
                    "n_with_manual_flow": len(flow),
                    "n_low_flow_samples": len(low),
                    "n_events": len(events),
                    "event_pct_of_flow_samples": safe_ratio(len(events) * 100, len(flow)),
                    "event_pct_of_low_flow_samples": safe_ratio(len(events) * 100, len(low)),
                    "median_pb_zn_ratio": group["pb_zn_ratio"].median(),
                }
            )
    return pd.DataFrame(rows)


def make_year_summary(samples: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for (system, year), group in samples.groupby(["system", "year"], sort=True):
        flow = group[group["manual_flow_percentile"].notna()]
        low = flow[flow["manual_low_flow"]]
        events = flow[flow["primary_event"]]
        rows.append(
            {
                "system": system,
                "year": int(year),
                "n_samples": len(group),
                "n_with_manual_flow": len(flow),
                "n_low_flow_samples": len(low),
                "n_events": len(events),
                "event_pct_of_low_flow_samples": safe_ratio(len(events) * 100, len(low)),
                "median_pb_ug_l": group["lead_value"].median(),
                "median_zn_ug_l": group["zinc_value"].median(),
                "median_pb_zn_ratio": group["pb_zn_ratio"].median(),
            }
        )
    return pd.DataFrame(rows)


def make_station_map(site_summary: pd.DataFrame) -> pd.DataFrame:
    columns = [
        "system",
        "station_id",
        "station_name",
        "station_type",
        "easting",
        "northing",
        "manual_nrfa_id",
        "manual_nrfa_name",
        "manual_decision",
        "manual_match_distance_km",
    ]
    return site_summary[columns].copy()


def plot_flow_ratio(samples: pd.DataFrame, threshold: float) -> Path:
    fig, ax = plt.subplots(figsize=(8.2, 5.2))
    for system in SYSTEM_ORDER:
        group = samples[
            samples["system"].eq(system) & samples["manual_flow_percentile"].notna()
        ]
        ax.scatter(
            group["manual_flow_percentile"],
            group["pb_zn_ratio"],
            s=np.where(group["primary_event"], 42, 24),
            color=PALETTE[system],
            alpha=np.where(group["primary_event"], 0.95, 0.55),
            edgecolor=np.where(group["primary_event"], "#222222", "none"),
            linewidth=np.where(group["primary_event"], 0.45, 0),
            label=SYSTEM_SHORT[system],
        )
    ax.axvspan(0, LOW_FLOW_PERCENTILE, color="#DDE8ED", alpha=0.55, zorder=0)
    ax.axvline(LOW_FLOW_PERCENTILE, color="#4A6674", linestyle="--", linewidth=1)
    ax.axhline(threshold, color="#9E3F38", linestyle="--", linewidth=1)
    ax.set_yscale("log")
    ax.set_xlim(-0.02, 1.02)
    ax.set_xlabel("Manually matched flow percentile")
    ax.set_ylabel("Pb/Zn ratio (log scale)")
    ax.set_title("Pb/Zn ratio in relation to manually validated river flow")
    ax.text(
        LOW_FLOW_PERCENTILE / 2,
        ax.get_ylim()[1] / 1.5,
        "Low-flow domain",
        ha="center",
        va="top",
        fontsize=8,
        color="#405965",
    )
    ax.legend(frameon=False, ncol=2, loc="upper right")
    path = RESULTS / "7.6_Flow_Ratio_Systems.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_event_drivers(driver_summary: pd.DataFrame) -> Path:
    data = driver_summary[~driver_summary["system"].eq("All four systems")]
    pivot = (
        data.pivot(index="system", columns="driver_category", values="n_events")
        .reindex(index=SYSTEM_ORDER, columns=DRIVER_ORDER)
        .fillna(0)
    )
    fig, ax = plt.subplots(figsize=(8.2, 4.6))
    left = np.zeros(len(pivot))
    for category in DRIVER_ORDER:
        values = pivot[category].to_numpy(float)
        bars = ax.barh(
            [SYSTEM_SHORT[index] for index in pivot.index],
            values,
            left=left,
            color=DRIVER_COLOURS[category],
            label=category,
            height=0.62,
        )
        for bar, value in zip(bars, values, strict=False):
            if value > 0:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_y() + bar.get_height() / 2,
                    str(int(value)),
                    ha="center",
                    va="center",
                    color="white" if category != "Neither marginal threshold" else "#222222",
                    fontsize=8,
                    fontweight="bold",
                )
        left += values
    ax.set_xlabel("Number of low-flow, high-Pb/Zn events")
    ax.set_ylabel("")
    ax.set_title("Threshold-based composition of the candidate events")
    ax.invert_yaxis()
    ax.legend(frameon=False, ncol=2, loc="lower right")
    path = RESULTS / "7.6_Event_Drivers.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_event_vs_normal(site_summary: pd.DataFrame) -> Path:
    data = site_summary.copy()
    data["label"] = data["station_id"]
    panels = [
        ("event_to_normal_pb_fold", "Pb", "#B45145"),
        ("event_to_normal_zn_fold", "Zn", "#3D7FA0"),
        ("event_to_normal_ratio_fold", "Pb/Zn", "#2F765E"),
    ]
    fig, axes = plt.subplots(1, 3, figsize=(10.2, 5.2), sharey=True, sharex=True)
    finite_values = np.concatenate(
        [data[column].dropna().to_numpy(float) for column, _, _ in panels]
    )
    x_min = max(0.4, float(np.nanmin(finite_values)) * 0.88)
    x_max = min(3.0, float(np.nanmax(finite_values)) * 1.10)
    for ax, (column, title, colour) in zip(axes, panels, strict=False):
        values = data[column].astype(float)
        ax.scatter(values, data["label"], s=48, color=colour, edgecolor="white", linewidth=0.6)
        ax.axvline(1, color="#555555", linestyle="--", linewidth=1)
        ax.set_xscale("log")
        ax.set_xlim(x_min, x_max)
        ax.xaxis.set_major_locator(mticker.FixedLocator([0.5, 1.0, 2.0]))
        ax.xaxis.set_major_formatter(mticker.FixedFormatter(["0.5", "1", "2"]))
        ax.xaxis.set_minor_formatter(mticker.NullFormatter())
        ax.set_xlabel("Event / normal-flow median")
        ax.set_title(title)
        ax.grid(axis="y", visible=False)
    axes[0].set_ylabel("Station ID")
    fig.suptitle(
        "How event chemistry differs from normal/high-flow chemistry at each site",
        y=1.02,
        fontsize=12,
    )
    path = RESULTS / "7.6_Event_vs_Normal.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_site_flow_correlations(site_summary: pd.DataFrame) -> Path:
    data = site_summary.sort_values("spearman_ratio_vs_flow_rho").copy()
    significant = data["spearman_ratio_vs_flow_fdr_q"].lt(0.05)
    colours = np.where(data["spearman_ratio_vs_flow_rho"].lt(0), "#2F765E", "#B56B38")
    fig, ax = plt.subplots(figsize=(7.8, 5.0))
    ax.scatter(
        data["spearman_ratio_vs_flow_rho"],
        data["station_id"],
        s=np.where(significant, 78, 50),
        color=colours,
        edgecolor=np.where(significant, "#111111", "white"),
        linewidth=np.where(significant, 1.1, 0.6),
        zorder=3,
    )
    ax.axvline(0, color="#555555", linestyle="--", linewidth=1)
    ax.set_xlim(-1, 1)
    ax.set_xlabel("Spearman correlation: Pb/Zn ratio vs flow percentile")
    ax.set_ylabel("Station ID")
    ax.set_title("Flow-Pb/Zn relationships differ among the nine robust stations")
    ax.scatter([], [], s=52, color="#2F765E", edgecolor="white", label="Negative rho")
    ax.scatter([], [], s=52, color="#B56B38", edgecolor="white", label="Positive rho")
    ax.scatter([], [], s=78, facecolor="white", edgecolor="#111111", linewidth=1.1, label="FDR q < 0.05")
    ax.legend(frameon=False, loc="lower right")
    ax.grid(axis="y", visible=False)
    path = RESULTS / "7.6_Site_Flow_Correlations.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_chemistry_context(samples: pd.DataFrame) -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(10.0, 4.6))
    for ax, variable, label in [
        (axes[0], "ph", "pH"),
        (axes[1], "calcium", "Calcium (mg/L)"),
    ]:
        complete = samples.dropna(subset=[variable, "pb_zn_ratio"])
        for system in SYSTEM_ORDER:
            group = complete[complete["system"].eq(system)]
            ax.scatter(
                group[variable],
                group["pb_zn_ratio"],
                s=27,
                color=PALETTE[system],
                alpha=0.65,
                edgecolor="none",
                label=SYSTEM_SHORT[system],
            )
        n, rho, p_value = safe_spearman(complete[variable], complete["pb_zn_ratio"])
        ax.set_yscale("log")
        ax.set_xlabel(label)
        ax.set_ylabel("Pb/Zn ratio (log scale)")
        p_text = "NA" if pd.isna(p_value) else f"{p_value:.3g}"
        rho_text = "NA" if pd.isna(rho) else f"{rho:.2f}"
        ax.set_title(f"n={n}; Spearman rho={rho_text}, p={p_text}")
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, frameon=False, ncol=4, loc="lower center", bbox_to_anchor=(0.5, -0.03))
    fig.suptitle("Exploratory pH and calcium context for Pb/Zn ratios", y=1.02, fontsize=12)
    path = RESULTS / "7.6_pH_Ca_Context.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_time_series(samples: pd.DataFrame, threshold: float) -> Path:
    fig, axes = plt.subplots(2, 2, figsize=(10.2, 7.2), constrained_layout=True)
    for ax, system in zip(axes.flat, SYSTEM_ORDER, strict=False):
        group = samples[samples["system"].eq(system)].sort_values("sample_date")
        regular = group[~group["primary_event"]]
        events = group[group["primary_event"]]
        ax.scatter(
            regular["sample_date"],
            regular["pb_zn_ratio"],
            s=17,
            color=PALETTE[system],
            alpha=0.48,
            edgecolor="none",
        )
        ax.scatter(
            events["sample_date"],
            events["pb_zn_ratio"],
            s=38,
            color="#C33F36",
            edgecolor="#222222",
            linewidth=0.4,
            label="Primary event",
            zorder=3,
        )
        ax.axhline(threshold, color="#777777", linestyle="--", linewidth=0.8)
        ax.set_yscale("log")
        ax.set_ylabel("Pb/Zn")
        ax.set_xlabel("Sample date")
        ax.set_title(SYSTEM_SHORT[system], loc="left", fontsize=10, fontweight="bold")
        ax.grid(axis="x", alpha=0.2)
        ax.tick_params(axis="x", labelrotation=0, labelsize=8)
    axes.flat[0].legend(frameon=False, loc="upper right")
    fig.suptitle("Time distribution of uncensored Pb/Zn samples and primary events", fontsize=12)
    path = RESULTS / "7.6_Time_Series.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def format_excel(path: Path, sheets: dict[str, pd.DataFrame]) -> None:
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        for sheet_name, frame in sheets.items():
            clean = frame.copy()
            for column in clean.select_dtypes(include=["datetime64[ns]"]).columns:
                clean[column] = clean[column].dt.strftime("%Y-%m-%d %H:%M:%S")
            clean.to_excel(writer, index=False, sheet_name=sheet_name[:31])
            ws = writer.book[sheet_name[:31]]
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
            for cell in ws[1]:
                cell.fill = PatternFill("solid", fgColor="1F4E78")
                cell.font = Font(color="FFFFFF", bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            ws.row_dimensions[1].height = 34
            for index, column in enumerate(ws.columns, start=1):
                values = [str(cell.value) if cell.value is not None else "" for cell in column[:300]]
                width = min(max(max(map(len, values), default=0) + 2, 11), 44)
                ws.column_dimensions[get_column_letter(index)].width = width
                for cell in column[1:]:
                    cell.alignment = Alignment(vertical="top", wrap_text=False)
                    if isinstance(cell.value, float):
                        cell.number_format = "0.000"


def make_readme(thresholds: dict[str, float], samples: pd.DataFrame) -> pd.DataFrame:
    events = samples[samples["primary_event"]]
    rows = [
        ("Purpose", "Final site- and system-level analysis of the nine robust candidate stations."),
        ("Primary data", "Uncensored same-time Pb-Zn pairs from the 7.5 sensitivity analysis."),
        ("Flow method", "Manually reviewed NRFA gauge match; same-day daily flow percentile."),
        ("Low-flow definition", f"Flow percentile <= {thresholds['low_flow_percentile']:.2f}."),
        ("High-Pb definition", f"Pb >= {thresholds['lead_high_ug_l']:.1f} ug/L (global uncensored 75th percentile)."),
        ("Low-Zn definition", f"Zn <= {thresholds['zinc_low_ug_l']:.1f} ug/L (global uncensored 25th percentile)."),
        ("High-ratio definition", f"Pb/Zn >= {thresholds['high_pb_zn_ratio']:.6f} (global uncensored 90th percentile)."),
        ("Primary event", "Low flow and high Pb/Zn ratio on the same sampled day."),
        ("Strict event", "Low flow, high Pb and low Zn on the same sampled day."),
        ("Robust stations", str(samples["station_id"].nunique())),
        ("Mine systems", str(samples["system"].nunique())),
        ("Uncensored samples", str(len(samples))),
        ("Samples with manual flow", str(samples["manual_flow_percentile"].notna().sum())),
        ("Primary events", str(len(events))),
        ("Strict events", str(int(events["strict_event"].sum()))),
        ("Important limitation", "Candidate stations were selected using the same event definition; outputs are descriptive and screening-oriented, not causal estimates."),
        ("Flow association", "Across the nine robust stations, Pb/Zn and flow percentile had Spearman rho = 0.009 (p = 0.883); site-specific directions were heterogeneous."),
        ("Calcium limitation", "Calcium coverage is incomplete and calcium is not equivalent to measured hardness."),
    ]
    return pd.DataFrame(rows, columns=["item", "definition_or_value"])


def set_run_font(run, size: float = 11, bold: bool | None = None, chinese: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun" if chinese else "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(document: Document, title: str, chinese: bool) -> None:
    section = document.sections[0]
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun" if chinese else "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size in [("Title", 17), ("Heading 1", 15), ("Heading 2", 12)]:
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun" if chinese else "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("17365D")
        style.paragraph_format.keep_with_next = True
    document.core_properties.title = title
    document.core_properties.author = "Zhou Xuchen"


def add_paragraph(document: Document, text: str, chinese: bool, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(paragraph.add_run(bold_prefix), bold=True, chinese=chinese)
        set_run_font(paragraph.add_run(text[len(bold_prefix):]), chinese=chinese)
    else:
        set_run_font(paragraph.add_run(text), chinese=chinese)


def add_dataframe_table(
    document: Document,
    frame: pd.DataFrame,
    columns: list[str],
    labels: list[str],
    chinese: bool,
) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, label in enumerate(labels):
        cell = table.rows[0].cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(label)
        set_run_font(run, 8.5, True, chinese=chinese)
        cell._tc.get_or_add_tcPr().append(_cell_shading("D9EAF2"))
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            value = row[column]
            if pd.isna(value):
                display = "-"
            elif isinstance(value, (float, np.floating)):
                display = f"{value:.2f}"
            else:
                display = str(value)
            cells[index].text = ""
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_run_font(cells[index].paragraphs[0].add_run(display), 8.3, chinese=chinese)


def _cell_shading(fill: str) -> OxmlElement:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    return shading


def add_picture(document: Document, path: Path, caption: str, chinese: bool) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(14.7))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(caption), 9, chinese=chinese)


def strongest_system_statement(system_summary: pd.DataFrame) -> tuple[str, int, float]:
    row = system_summary.sort_values(
        ["n_low_flow_high_ratio_events", "event_pct_of_flow_samples"], ascending=False
    ).iloc[0]
    return str(row["system"]), int(row["n_low_flow_high_ratio_events"]), float(row["event_pct_of_flow_samples"])


def driver_counts(driver_summary: pd.DataFrame) -> dict[str, int]:
    overall = driver_summary[driver_summary["system"].eq("All four systems")]
    return {
        row["driver_category"]: int(row["n_events"])
        for _, row in overall.iterrows()
    }


def build_chinese_summary(
    samples: pd.DataFrame,
    site_summary: pd.DataFrame,
    system_summary: pd.DataFrame,
    driver_summary: pd.DataFrame,
    chemistry: pd.DataFrame,
    thresholds: dict[str, float],
    figures: dict[str, Path],
) -> Path:
    path = RESULTS / "7.6_Final_Analysis_Summary_CN.docx"
    document = Document()
    configure_document(document, "最终分析总结", chinese=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("最终分析总结"), 17, True, chinese=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        subtitle.add_run("九个稳健候选站与四个矿区系统的 Pb/Zn-流量分析"),
        11,
        chinese=True,
    )

    total = len(samples)
    flow_n = int(samples["manual_flow_percentile"].notna().sum())
    events = samples[samples["primary_event"]]
    strict_n = int(events["strict_event"].sum())
    strongest, strongest_n, strongest_pct = strongest_system_statement(system_summary)
    counts = driver_counts(driver_summary)
    ph_row = chemistry[
        chemistry["analysis_group"].eq("All four systems") & chemistry["variable"].eq("ph")
    ].iloc[0]
    ca_row = chemistry[
        chemistry["analysis_group"].eq("All four systems") & chemistry["variable"].eq("calcium")
    ].iloc[0]

    document.add_heading("1. 分析目的与最终方法", level=1)
    add_paragraph(
        document,
        "本阶段用于完成论文写作前的最终分析。分析对象为经过检出限敏感性、同日/3日/7日流量窗口以及人工流量站核对后均保持稳定的9个候选水质站，并按空间位置与矿区关系归并为4个系统：Esgair Mwyn/Nant y Garw、Nant y Watcyn、Esgair Hir/Eastern Ffraith和Wemyss。导师已经认可保留自动匹配信息并用人工核对验证的方法；因此最终结果采用人工确认的NRFA流量站及同日流量百分位。",
        chinese=True,
    )
    add_paragraph(
        document,
        f"主分析只使用未删失的同一采样时间Pb-Zn配对。低流量定义为流量百分位不高于{thresholds['low_flow_percentile']:.2f}；高Pb为不低于{thresholds['lead_high_ug_l']:.0f} µg/L；低Zn为不高于{thresholds['zinc_low_ug_l']:.0f} µg/L；高Pb/Zn为不低于{thresholds['high_pb_zn_ratio']:.3f}。主要事件指同一天同时满足低流量和高Pb/Zn，严格事件还要求高Pb与低Zn同时出现。",
        chinese=True,
    )

    document.add_heading("2. 数据覆盖与主要结果", level=1)
    add_paragraph(
        document,
        f"九个站共有{total}个未删失Pb-Zn样本，其中{flow_n}个能够连接人工确认的同日流量。共识别{len(events)}个低流量-高Pb/Zn主要事件，其中{strict_n}个同时达到高Pb和低Zn的严格条件。所有九个站在此前各项敏感性检验中均保持候选状态，因此这里的分析用于描述这些系统的化学特征与优先级，而不是重新筛选候选站。",
        chinese=True,
    )
    add_paragraph(
        document,
        f"按事件数量计，{strongest}最突出，共{strongest_n}个事件，占该系统有流量样本的{strongest_pct:.1f}%。但该系统包含5个站，采样量也最大，因此不同系统之间不能只按事件总数直接比较；系统表同时提供事件占全部流量样本和低流量样本的比例。",
        chinese=True,
    )

    display_system = system_summary[
        [
            "system",
            "n_stations",
            "n_with_manual_flow",
            "n_low_flow_high_ratio_events",
            "n_strict_high_pb_low_zn_events",
            "event_pct_of_flow_samples",
            "event_station_normalised_pb_median",
            "event_station_normalised_zn_median",
            "event_station_normalised_ratio_median",
        ]
    ]
    add_dataframe_table(
        document,
        display_system,
        list(display_system.columns),
        ["矿区系统", "站点数", "流量样本", "主要事件", "严格事件", "事件率(%)", "Pb倍数", "Zn倍数", "比值倍数"],
        chinese=True,
    )
    add_picture(document, figures["flow_ratio"], "图1  人工匹配流量百分位与Pb/Zn比值", chinese=True)
    add_paragraph(
        document,
        "九个站合并后，Pb/Zn与流量百分位的Spearman相关系数为0.009（p=0.883），即没有总体单调关系。站点层面经Benjamini-Hochberg多重检验校正后，S83020呈显著负相关（ρ=-0.619），符合流量越低比值越高；S83018呈显著正相关（ρ=0.430），方向与假设相反。其余七站没有显著单调关系。",
        chinese=True,
    )
    add_picture(document, figures["site_correlations"], "图2  九个稳健站的Pb/Zn-流量相关方向", chinese=True)

    document.add_heading("3. 高Pb/Zn事件由Pb还是Zn驱动", level=1)
    add_paragraph(
        document,
        "为了回答‘高Pb、低Zn’是否为主要机制，每个主要事件按预先确定的Pb和Zn阈值分成四类。该分类不引入新的数据驱动阈值，因此可以清楚区分严格的高Pb-低Zn事件、单独高Pb、单独低Zn，以及两项浓度都没有越过边际阈值但比值仍然较高的相对组成变化。",
        chinese=True,
    )
    add_paragraph(
        document,
        f"在{len(events)}个主要事件中，高Pb与低Zn同时出现{counts['High Pb and low Zn']}次，仅高Pb出现{counts['High Pb only']}次，仅低Zn出现{counts['Low Zn only']}次，另有{counts['Neither marginal threshold']}次没有单独越过Pb或Zn阈值。由此可见，‘低流量-高Pb/Zn’并不等同于每次都出现绝对高Pb和绝对低Zn；系统解释必须同时查看Pb、Zn及其相对于正常/高流量样本的倍数变化。",
        chinese=True,
    )
    add_picture(document, figures["drivers"], "图3  主要事件的Pb/Zn阈值组成", chinese=True)
    add_picture(document, figures["folds"], "图4  各站事件期与正常/高流量期中位数之比", chinese=True)
    add_paragraph(
        document,
        "站点内比较进一步说明，共现不等于低流量放大。S83020的事件期Pb中位数是正常/高流量期的2.11倍，Zn约为1.01倍，Pb/Zn为2.33倍，是最清楚的低流量Pb富集信号；S83017的比值为1.40倍，属于提示性结果。S83018、S83019、S83021和S6320066的事件期比值反而略低于各自正常/高流量期，说明它们主要是本身长期具有较高比值，低流量日只是与高比值重叠。",
        chinese=True,
    )

    document.add_heading("4. pH、Ca、季节与时间背景", level=1)
    add_paragraph(
        document,
        f"pH在九个稳健站中有{int(ph_row['n_complete_pairs'])}个完整配对，其与Pb/Zn的总体Spearman相关系数为{ph_row['spearman_rho_with_pb_zn_ratio']:.2f}。Ca只有{int(ca_row['n_complete_pairs'])}个完整配对，总体相关系数为{ca_row['spearman_rho_with_pb_zn_ratio']:.2f}。这些关系只作探索性描述，因为站点组成、采样年份和采样频率不均衡会造成混杂，而且Ca不能替代实验测得的硬度。原始六指标数据中的硬度没有有效记录，因此本阶段不能检验硬度效应。",
        chinese=True,
    )
    add_paragraph(
        document,
        "季节表分别给出各系统每季的样本数、低流量样本数、事件数，以及事件占低流量样本的比例。时间序列图显示事件出现的具体年份和采样空档。由于采样并非等时间间隔，季节和年度结果应理解为采样记录中的分布，而不是自然发生频率或趋势检验。",
        chinese=True,
    )
    add_picture(document, figures["chemistry"], "图5  Pb/Zn与pH及Ca的探索性关系", chinese=True)
    add_picture(document, figures["time"], "图6  四个矿区系统的Pb/Zn时间分布", chinese=True)

    document.add_heading("5. 可写入论文的结论", level=1)
    conclusions = [
        "人工核对后的流量匹配没有否定九个稳健候选站，说明经验证的简化匹配流程足以支持区域筛选。",
        f"九个站在{flow_n}个有流量的未删失样本中出现{len(events)}个低流量-高Pb/Zn事件，表明这些站值得作为后续现场核查和修复优先级评估对象。",
        "严格的高Pb-低Zn事件只有4个，全部位于Esgair Mwyn/Nant y Garw系统，其中S83018有3个、S83021有1个。",
        "S83020是低流量放大假设证据最强的站点；S83017为提示性结果；S83018的显著关系方向与假设相反。",
        "只有一部分主要事件同时达到绝对高Pb和绝对低Zn，因此Pb/Zn升高可能由Pb升高、Zn降低或两者的相对变化共同形成，不能只用比值推断单一机制。",
        "合并九站没有发现Pb/Zn随流量下降而系统升高，因此论文必须把事件筛选、站点内关联和因果解释明确区分。",
        "pH可用于辅助解释金属迁移，而Ca覆盖不足、硬度完全缺失，限制了水化学缓冲能力的评价。",
        "候选站是通过同一事件定义筛选得到的，最终结果应表述为稳健的筛选证据和描述性系统比较，而不是流量导致金属变化的因果证明。",
    ]
    for item in conclusions:
        paragraph = document.add_paragraph(style="List Number")
        set_run_font(paragraph.add_run(item), chinese=True)

    document.add_heading("6. 输出文件说明", level=1)
    add_paragraph(
        document,
        "7.6_Site_Profiles.xlsx包含方法说明、站点汇总、系统汇总、事件驱动分类、全部事件、pH/Ca相关、季节汇总、年度汇总、样本级数据以及站点-流量站对应表。CSV文件可直接由Python/R继续处理；PNG图片可用于论文Results章节，图注需要在论文中说明阈值、人工流量匹配和未删失样本限制。",
        chinese=True,
    )
    document.save(path)
    return path


def build_english_summary(
    samples: pd.DataFrame,
    system_summary: pd.DataFrame,
    driver_summary: pd.DataFrame,
    chemistry: pd.DataFrame,
    thresholds: dict[str, float],
    figures: dict[str, Path],
) -> Path:
    path = RESULTS / "7.6_Final_Analysis_Summary_EN.docx"
    document = Document()
    configure_document(document, "Final Analysis Summary", chinese=False)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("Final Analysis Summary"), 17, True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("Nine robust stations grouped into four mine systems"), 11)

    flow_n = int(samples["manual_flow_percentile"].notna().sum())
    events = samples[samples["primary_event"]]
    strict_n = int(events["strict_event"].sum())
    counts = driver_counts(driver_summary)
    ph_row = chemistry[
        chemistry["analysis_group"].eq("All four systems") & chemistry["variable"].eq("ph")
    ].iloc[0]
    ca_row = chemistry[
        chemistry["analysis_group"].eq("All four systems") & chemistry["variable"].eq("calcium")
    ].iloc[0]

    document.add_heading("1. Scope and primary method", level=1)
    add_paragraph(
        document,
        "This final analysis uses the nine stations that remained candidates under the censoring, flow-window and manual flow-matching checks. They were grouped into four mine systems: Esgair Mwyn/Nant y Garw, Nant y Watcyn, Esgair Hir/Eastern Ffraith and Wemyss. The primary results use the manually reviewed NRFA gauge and same-day daily-flow percentile.",
        chinese=False,
    )
    add_paragraph(
        document,
        f"Only uncensored, same-time Pb-Zn pairs were retained. Low flow was defined as a flow percentile <= {thresholds['low_flow_percentile']:.2f}; high Pb as >= {thresholds['lead_high_ug_l']:.0f} ug/L; low Zn as <= {thresholds['zinc_low_ug_l']:.0f} ug/L; and high Pb/Zn as >= {thresholds['high_pb_zn_ratio']:.3f}. A primary event combined low flow and high Pb/Zn on the sampled day. A strict event additionally combined high Pb and low Zn.",
        chinese=False,
    )

    document.add_heading("2. Main findings", level=1)
    add_paragraph(
        document,
        f"The nine sites contained {len(samples)} uncensored Pb-Zn pairs, of which {flow_n} had manually matched flow. There were {len(events)} primary low-flow, high-Pb/Zn events and {strict_n} strict high-Pb, low-Zn events.",
        chinese=False,
    )
    display_system = system_summary[
        [
            "system",
            "n_stations",
            "n_with_manual_flow",
            "n_low_flow_high_ratio_events",
            "n_strict_high_pb_low_zn_events",
            "event_pct_of_flow_samples",
            "event_station_normalised_pb_median",
            "event_station_normalised_zn_median",
            "event_station_normalised_ratio_median",
        ]
    ]
    add_dataframe_table(
        document,
        display_system,
        list(display_system.columns),
        ["Mine system", "Sites", "Flow n", "Events", "Strict", "Event %", "Pb fold", "Zn fold", "Ratio fold"],
        chinese=False,
    )
    add_picture(document, figures["flow_ratio"], "Figure 1. Manually matched flow percentile and Pb/Zn ratio.", chinese=False)
    add_paragraph(
        document,
        "Across all nine stations, Pb/Zn and flow percentile were not monotonically associated (Spearman rho = 0.009, p = 0.883). After Benjamini-Hochberg correction across the nine site tests, S83020 showed a significant negative association (rho = -0.619), consistent with ratios rising as flow fell, whereas S83018 showed a significant positive association (rho = 0.430), opposite to the hypothesis. The other seven site correlations were not significant.",
        chinese=False,
    )
    add_picture(document, figures["site_correlations"], "Figure 2. Site-specific directions of the Pb/Zn-flow relationship.", chinese=False)

    document.add_heading("3. Event composition", level=1)
    add_paragraph(
        document,
        f"Among the {len(events)} primary events, {counts['High Pb and low Zn']} combined high Pb and low Zn, {counts['High Pb only']} had high Pb only, {counts['Low Zn only']} had low Zn only, and {counts['Neither marginal threshold']} crossed neither marginal threshold. High Pb/Zn under low flow therefore did not always represent simultaneous absolute high Pb and low Zn. The component concentrations and their event-to-normal-flow fold changes must be interpreted with the ratio.",
        chinese=False,
    )
    add_picture(document, figures["drivers"], "Figure 3. Threshold-based composition of primary events.", chinese=False)
    add_picture(document, figures["folds"], "Figure 4. Event-to-normal/high-flow median chemistry at each station.", chinese=False)
    add_paragraph(
        document,
        "Within-site normalisation separates chronic high-ratio sites from low-flow amplification. At S83020, event Pb was 2.11 times the normal/high-flow median, Zn was 1.01 times, and Pb/Zn was 2.33 times, giving the clearest signal of low-flow Pb enrichment. S83017 was suggestive (ratio fold 1.40). Event ratios at S83018, S83019, S83021 and S6320066 were slightly below their own normal/high-flow medians, indicating co-occurrence at chronically high-ratio sites rather than low-flow amplification.",
        chinese=False,
    )

    document.add_heading("4. Supporting chemistry and time", level=1)
    add_paragraph(
        document,
        f"pH had {int(ph_row['n_complete_pairs'])} complete pairs and an overall Spearman correlation of {ph_row['spearman_rho_with_pb_zn_ratio']:.2f} with Pb/Zn. Calcium had only {int(ca_row['n_complete_pairs'])} complete pairs and an overall correlation of {ca_row['spearman_rho_with_pb_zn_ratio']:.2f}. These are exploratory associations because stations, years and sampling effort are unbalanced. Calcium is not equivalent to measured hardness, and hardness could not be analysed because no valid hardness results were present.",
        chinese=False,
    )
    add_picture(document, figures["chemistry"], "Figure 5. Exploratory pH and calcium context.", chinese=False)
    add_picture(document, figures["time"], "Figure 6. Time distribution of Pb/Zn samples and primary events.", chinese=False)

    document.add_heading("5. Thesis-ready interpretation", level=1)
    conclusions = [
        "Manual review did not remove the nine robust candidates, supporting the validated matching workflow as a regional screening method.",
        f"The {len(events)} primary events identify systems suitable for field verification and remediation-priority assessment.",
        "All four strict high-Pb, low-Zn events occurred in the Esgair Mwyn/Nant y Garw system: three at S83018 and one at S83021.",
        "S83020 provided the strongest evidence for low-flow amplification, S83017 was suggestive, and S83018 had a significant association in the opposite direction.",
        "Only a subset of events met both absolute high-Pb and low-Zn criteria; elevated ratios can reflect Pb enrichment, Zn depletion or relative changes in both components.",
        "The pooled nine-site analysis did not show Pb/Zn systematically increasing as flow declined; event screening, within-site association and causal interpretation must remain separate.",
        "The results are robust screening and descriptive-comparison evidence, not proof that low river flow caused the observed metal chemistry.",
        "Unequal sampling, incomplete calcium, absent hardness and daily rather than sub-daily flow remain the principal limitations.",
    ]
    for item in conclusions:
        paragraph = document.add_paragraph(style="List Number")
        set_run_font(paragraph.add_run(item))

    document.save(path)
    return path


def export_outputs(
    samples: pd.DataFrame,
    site_summary: pd.DataFrame,
    system_summary: pd.DataFrame,
    drivers: pd.DataFrame,
    chemistry: pd.DataFrame,
    seasons: pd.DataFrame,
    years: pd.DataFrame,
    station_map: pd.DataFrame,
    thresholds: dict[str, float],
) -> dict[str, pd.DataFrame]:
    events = samples[samples["primary_event"]].copy()
    readme = make_readme(thresholds, samples)
    tables = {
        "README": readme,
        "Site_Summary": site_summary,
        "System_Summary": system_summary,
        "Event_Drivers": drivers,
        "Event_Samples": events,
        "Chemistry_Context": chemistry,
        "Season_Summary": seasons,
        "Year_Summary": years,
        "Sample_Data": samples,
        "Station_Flow_Map": station_map,
    }
    format_excel(RESULTS / "7.6_Site_Profiles.xlsx", tables)
    csv_outputs = {
        "7.6_Site_Summary.csv": site_summary,
        "7.6_System_Summary.csv": system_summary,
        "7.6_Event_Drivers.csv": drivers,
        "7.6_Event_Samples.csv": events,
        "7.6_Chemistry_Context.csv": chemistry,
        "7.6_Season_Summary.csv": seasons,
        "7.6_Year_Summary.csv": years,
        "7.6_Sample_Data.csv": samples,
        "7.6_Station_Flow_Map.csv": station_map,
    }
    for filename, frame in csv_outputs.items():
        frame.to_csv(RESULTS / filename, index=False)
    return tables


def audit_outputs(
    samples: pd.DataFrame,
    site_summary: pd.DataFrame,
    system_summary: pd.DataFrame,
    drivers: pd.DataFrame,
) -> None:
    assert samples["station_id"].nunique() == 9
    assert samples["system"].nunique() == 4
    assert set(samples["station_id"]) == set(SYSTEMS)
    assert len(samples) == int(site_summary["n_uncensored_pbzn"].sum())
    assert int(samples["primary_event"].sum()) == int(
        system_summary["n_low_flow_high_ratio_events"].sum()
    )
    overall_driver_count = int(
        drivers.loc[drivers["system"].eq("All four systems"), "n_events"].sum()
    )
    assert overall_driver_count == int(samples["primary_event"].sum())
    assert not samples.loc[samples["primary_event"], "manual_flow_percentile"].gt(
        LOW_FLOW_PERCENTILE
    ).any()
    assert not samples.loc[samples["primary_event"], "high_pb_zn_ratio"].eq(False).any()


def main() -> None:
    RESULTS.mkdir(parents=True, exist_ok=True)
    setup_plot_style()
    samples, robust, thresholds = load_analysis_samples()
    site_summary = make_site_summary(samples, robust)
    system_summary = make_system_summary(samples)
    drivers = make_driver_summary(samples)
    chemistry = make_chemistry_context(samples)
    seasons = make_season_summary(samples)
    years = make_year_summary(samples)
    station_map = make_station_map(site_summary)
    audit_outputs(samples, site_summary, system_summary, drivers)
    export_outputs(
        samples,
        site_summary,
        system_summary,
        drivers,
        chemistry,
        seasons,
        years,
        station_map,
        thresholds,
    )
    figures = {
        "flow_ratio": plot_flow_ratio(samples, thresholds["high_pb_zn_ratio"]),
        "site_correlations": plot_site_flow_correlations(site_summary),
        "drivers": plot_event_drivers(drivers),
        "folds": plot_event_vs_normal(site_summary),
        "chemistry": plot_chemistry_context(samples),
        "time": plot_time_series(samples, thresholds["high_pb_zn_ratio"]),
    }
    cn = build_chinese_summary(
        samples,
        site_summary,
        system_summary,
        drivers,
        chemistry,
        thresholds,
        figures,
    )
    en = build_english_summary(
        samples,
        system_summary,
        drivers,
        chemistry,
        thresholds,
        figures,
    )
    print(
        {
            "output_directory": str(RESULTS),
            "n_samples": len(samples),
            "n_manual_flow": int(samples["manual_flow_percentile"].notna().sum()),
            "n_events": int(samples["primary_event"].sum()),
            "n_strict_events": int(samples["strict_event"].sum()),
            "chinese_summary": str(cn),
            "english_summary": str(en),
        }
    )


if __name__ == "__main__":
    main()
